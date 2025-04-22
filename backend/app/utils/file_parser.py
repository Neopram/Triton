import os
import io
import uuid
import json
import logging
import concurrent.futures
from datetime import datetime
from typing import Dict, List, Union, Optional, Tuple, BinaryIO, Any, Callable
from pathlib import Path

import pdfplumber
import pandas as pd
import numpy as np
import magic
import chardet
import xlrd
import openpyxl
import docx
import pytesseract
from PIL import Image
from pdf2image import convert_from_path, convert_from_bytes
from openpyxl import load_workbook
from fastnumbers import fast_float, fast_int
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.core.logging import file_logger
from app.models.document import DocumentMetadata, ProcessingLog
from app.services.storage_service import upload_to_storage, get_from_storage
from app.utils.text_processor import clean_text, summarize_text, extract_keywords
from app.utils.validation import validate_column_names, validate_data_types
from app.utils.security import scan_file_for_threats, sanitize_content
from app.exceptions.file_exceptions import (
    FileParsingError,
    UnsupportedFileTypeError,
    CorruptFileError,
    FileSecurityError
)

# Configure logging
logger = logging.getLogger(__name__)

# Supported file types with their MIME types
SUPPORTED_FILE_TYPES = {
    "pdf": ["application/pdf"],
    "csv": ["text/csv", "application/csv", "application/vnd.ms-excel"],
    "excel": ["application/vnd.ms-excel", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"],
    "txt": ["text/plain"],
    "json": ["application/json"],
    "docx": ["application/vnd.openxmlformats-officedocument.wordprocessingml.document"],
    "image": ["image/jpeg", "image/png", "image/tiff", "image/gif"]
}

# Enterprise-level configuration
FILE_PROCESSING_CONFIG = {
    "max_file_size_mb": settings.MAX_FILE_SIZE_MB,
    "max_pages_default": settings.MAX_PDF_PAGES,
    "ocr_enabled": settings.OCR_ENABLED,
    "security_scan_enabled": settings.SECURITY_SCAN_ENABLED,
    "parallel_processing": settings.PARALLEL_PROCESSING_ENABLED,
    "max_workers": settings.MAX_PROCESSING_WORKERS,
    "timeout_seconds": settings.FILE_PROCESSING_TIMEOUT,
    "extract_metadata": settings.EXTRACT_FILE_METADATA,
    "temp_directory": settings.TEMP_FILE_DIRECTORY
}


class EnterpriseFileParser:
    """
    Enterprise-grade file parsing utility with comprehensive format support,
    security features, and performance optimizations.
    """
    
    def __init__(self, db: Optional[AsyncSession] = None):
        """
        Initialize the parser with optional database session for logging.
        
        Args:
            db: Optional database session for persistent logging
        """
        self.db = db
        self.temp_dir = Path(FILE_PROCESSING_CONFIG["temp_directory"])
        
        # Ensure temp directory exists
        if not self.temp_dir.exists():
            self.temp_dir.mkdir(parents=True, exist_ok=True)
        
        self.processing_stats = {
            "start_time": None,
            "end_time": None,
            "file_size": 0,
            "processing_duration": 0,
            "error_count": 0,
            "warnings": []
        }
    
    async def parse_file(
        self, 
        file_path: Union[str, Path, BinaryIO],
        file_type: Optional[str] = None,
        options: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Smart file parsing with automatic type detection and comprehensive processing.
        
        Args:
            file_path: Path to file, Path object, or file-like object
            file_type: Optional explicit file type override
            options: Optional processing parameters
            
        Returns:
            Dictionary with parsed content and metadata
        """
        self.processing_stats["start_time"] = datetime.utcnow()
        
        # Initialize options with defaults if not provided
        if options is None:
            options = {}
        
        # Generate a unique processing ID
        processing_id = str(uuid.uuid4())
        file_logger.info(f"Starting file processing with ID: {processing_id}")
        
        try:
            # Handle file input (path string, Path object, or file-like object)
            temp_file_path, should_delete = await self._prepare_input_file(file_path)
            
            # Get file size for logging
            file_size = os.path.getsize(temp_file_path)
            self.processing_stats["file_size"] = file_size
            
            # Check file size limit
            max_size = FILE_PROCESSING_CONFIG["max_file_size_mb"] * 1024 * 1024
            if file_size > max_size:
                raise FileParsingError(f"File exceeds maximum size limit of {FILE_PROCESSING_CONFIG['max_file_size_mb']}MB")
            
            # Security scan if enabled
            if FILE_PROCESSING_CONFIG["security_scan_enabled"]:
                scan_result = scan_file_for_threats(temp_file_path)
                if not scan_result["safe"]:
                    raise FileSecurityError(f"Security scan failed: {scan_result['reason']}")
            
            # Detect file type if not provided
            detected_mime_type = self._detect_mime_type(temp_file_path)
            if not file_type:
                file_type = self._determine_file_type(detected_mime_type)
            
            file_logger.info(f"Processing {file_type} file ({file_size/1024:.2f}KB): {temp_file_path}")
            
            # Extract file metadata if configured
            metadata = {}
            if FILE_PROCESSING_CONFIG["extract_metadata"]:
                metadata = self._extract_file_metadata(temp_file_path, file_type, detected_mime_type)
            
            # Parse content based on file type
            parsing_method = self._get_parsing_method(file_type)
            
            # Process with appropriate parser
            content = await parsing_method(temp_file_path, options)
            
            # Clean up temporary file if needed
            if should_delete:
                os.remove(temp_file_path)
            
            # Prepare result
            self.processing_stats["end_time"] = datetime.utcnow()
            processing_duration = (self.processing_stats["end_time"] - self.processing_stats["start_time"]).total_seconds()
            self.processing_stats["processing_duration"] = processing_duration
            
            result = {
                "processing_id": processing_id,
                "content": content,
                "metadata": metadata,
                "file_type": file_type,
                "mime_type": detected_mime_type,
                "processing_stats": {
                    "duration_seconds": processing_duration,
                    "file_size_bytes": file_size,
                    "processed_at": self.processing_stats["end_time"].isoformat(),
                    "warnings": self.processing_stats["warnings"]
                }
            }
            
            # Log result summary
            file_logger.info(
                f"File processed successfully in {processing_duration:.2f}s: "
                f"ID={processing_id}, Type={file_type}, Size={file_size/1024:.2f}KB"
            )
            
            # Record in database if session provided
            if self.db:
                await self._record_processing_log(processing_id, True, result["processing_stats"])
            
            return result
            
        except Exception as e:
            self.processing_stats["error_count"] += 1
            self.processing_stats["end_time"] = datetime.utcnow()
            processing_duration = (self.processing_stats["end_time"] - self.processing_stats["start_time"]).total_seconds()
            
            error_info = {
                "processing_id": processing_id,
                "error": str(e),
                "error_type": type(e).__name__,
                "duration_seconds": processing_duration
            }
            
            file_logger.error(
                f"File processing failed in {processing_duration:.2f}s: "
                f"ID={processing_id}, Error={str(e)}"
            )
            
            # Record error in database if session provided
            if self.db:
                await self._record_processing_log(processing_id, False, error_info)
            
            # Re-raise as appropriate exception type
            if isinstance(e, (FileParsingError, UnsupportedFileTypeError, CorruptFileError, FileSecurityError)):
                raise
            else:
                raise FileParsingError(f"Error processing file: {str(e)}") from e
    
    async def parse_pdf(
        self, 
        file_path: Union[str, Path],
        options: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Advanced PDF parsing with text extraction, OCR capabilities, and structural analysis.
        
        Args:
            file_path: Path to PDF file
            options: Dictionary of options including:
                - page_range: Tuple[int, int] - Range of pages to process
                - extract_tables: bool - Whether to extract tables
                - ocr_fallback: bool - Whether to use OCR if text extraction fails
                - extract_images: bool - Whether to extract images
                - password: str - PDF password if encrypted
                
        Returns:
            Dictionary with PDF content and structure
        """
        if options is None:
            options = {}
        
        # Default options
        page_range = options.get("page_range", None)
        extract_tables = options.get("extract_tables", True)
        ocr_fallback = options.get("ocr_fallback", FILE_PROCESSING_CONFIG["ocr_enabled"])
        extract_images = options.get("extract_images", False)
        password = options.get("password", None)
        
        try:
            pdf_content = {
                "text": "",
                "pages": [],
                "tables": [],
                "images": [],
                "structure": {}
            }
            
            # Open PDF with password if provided
            pdf_kwargs = {"password": password} if password else {}
            
            with pdfplumber.open(file_path, **pdf_kwargs) as pdf:
                total_pages = len(pdf.pages)
                
                # Determine page range to process
                if page_range:
                    start_page, end_page = page_range
                    start_page = max(0, start_page)
                    end_page = min(total_pages, end_page)
                else:
                    start_page = 0
                    end_page = min(total_pages, FILE_PROCESSING_CONFIG["max_pages_default"])
                
                # Extract document metadata
                pdf_content["structure"] = {
                    "total_pages": total_pages,
                    "processed_pages": end_page - start_page,
                    "metadata": pdf.metadata
                }
                
                # Process pages in parallel if configured
                if FILE_PROCESSING_CONFIG["parallel_processing"] and total_pages > 1:
                    pages_to_process = list(range(start_page, end_page))
                    with concurrent.futures.ThreadPoolExecutor(max_workers=FILE_PROCESSING_CONFIG["max_workers"]) as executor:
                        future_to_page = {
                            executor.submit(
                                self._process_pdf_page, 
                                pdf.pages[i], 
                                extract_tables, 
                                ocr_fallback
                            ): i for i in pages_to_process
                        }
                        
                        page_results = []
                        for future in concurrent.futures.as_completed(future_to_page):
                            page_index = future_to_page[future]
                            try:
                                page_result = future.result()
                                page_result["page_number"] = page_index + 1
                                page_results.append((page_index, page_result))
                            except Exception as e:
                                self.processing_stats["warnings"].append(
                                    f"Error processing page {page_index + 1}: {str(e)}"
                                )
                        
                        # Sort results by page number
                        page_results.sort(key=lambda x: x[0])
                        pdf_content["pages"] = [result for _, result in page_results]
                else:
                    # Process pages sequentially
                    for page_num in range(start_page, end_page):
                        page = pdf.pages[page_num]
                        page_result = self._process_pdf_page(page, extract_tables, ocr_fallback)
                        page_result["page_number"] = page_num + 1
                        pdf_content["pages"].append(page_result)
                
                # Extract all text
                pdf_content["text"] = "\n\n".join(page["text"] for page in pdf_content["pages"])
                
                # Combine tables from all pages
                pdf_content["tables"] = [
                    table for page in pdf_content["pages"] 
                    for table in page.get("tables", [])
                ]
                
                # Extract images if requested
                if extract_images:
                    pdf_content["images"] = self._extract_pdf_images(
                        file_path, 
                        range(start_page, end_page)
                    )
            
            return pdf_content
            
        except Exception as e:
            raise FileParsingError(f"Error parsing PDF: {str(e)}")
    
    async def parse_csv(
        self, 
        file_path: Union[str, Path],
        options: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Advanced CSV parsing with comprehensive options and data validation.
        
        Args:
            file_path: Path to CSV file
            options: Dictionary of options including:
                - encoding: str - File encoding
                - delimiter: str - Field delimiter
                - quotechar: str - Character for quotes
                - header: Union[int, List[str]] - Header row index or custom headers
                - skip_rows: int - Number of rows to skip
                - parse_dates: List[str] - Columns to parse as dates
                - thousands: str - Thousands separator
                - convert_numbers: bool - Auto-convert numeric strings
                - validate_schema: Dict - Schema for validation
                
        Returns:
            Dictionary with parsed CSV data and metadata
        """
        if options is None:
            options = {}
        
        try:
            # Detect encoding if not specified
            encoding = options.get("encoding", None)
            if not encoding:
                with open(file_path, 'rb') as f:
                    result = chardet.detect(f.read(10000))
                    encoding = result['encoding']
            
            # Prepare pandas read_csv options
            pandas_options = {
                "encoding": encoding,
                "delimiter": options.get("delimiter", ","),
                "quotechar": options.get("quotechar", '"'),
                "skiprows": options.get("skip_rows", 0),
                "thousands": options.get("thousands", ","),
                "na_values": options.get("na_values", ["", "NA", "N/A", "null", "NULL", "None"]),
                "low_memory": False
            }
            
            # Handle custom header
            if "header" in options:
                if isinstance(options["header"], int):
                    pandas_options["header"] = options["header"]
                else:
                    pandas_options["header"] = None
                    pandas_options["names"] = options["header"]
            
            # Parse date columns if specified
            if "parse_dates" in options:
                pandas_options["parse_dates"] = options["parse_dates"]
            
            # Read the CSV
            df = pd.read_csv(file_path, **pandas_options)
            
            # Convert numbers if requested
            if options.get("convert_numbers", True):
                for col in df.columns:
                    # Try to convert to numeric, but only if most values in the column appear numeric
                    if df[col].dtype == 'object':
                        # Sample values to check if they're numbers
                        sample = df[col].dropna().head(100)
                        if sample.empty:
                            continue
                            
                        # Try converting the sample to see if more than 90% convert successfully
                        numeric_count = 0
                        for val in sample:
                            try:
                                float(val)
                                numeric_count += 1
                            except (ValueError, TypeError):
                                pass
                        
                        if numeric_count / len(sample) > 0.9:
                            # Apply fast numeric conversion
                            df[col] = df[col].apply(
                                lambda x: fast_float(x) if x is not None and x != "" else None
                            )
            
            # Validate against schema if provided
            validation_results = {}
            if "validate_schema" in options:
                schema = options["validate_schema"]
                validation_results = validate_data_types(df, schema)
                
                # Validate column names if schema includes them
                if "columns" in schema:
                    name_validation = validate_column_names(
                        df.columns.tolist(), 
                        schema["columns"]
                    )
                    validation_results["column_names"] = name_validation
            
            # Build result
            result = {
                "data": df.to_dict(orient="records"),
                "columns": df.columns.tolist(),
                "row_count": len(df),
                "column_count": len(df.columns),
                "sample": df.head(5).to_dict(orient="records"),
                "summary": {
                    "numeric_columns": df.select_dtypes(include=[np.number]).columns.tolist(),
                    "date_columns": df.select_dtypes(include=["datetime"]).columns.tolist(),
                    "string_columns": df.select_dtypes(include=["object"]).columns.tolist(),
                    "missing_values": df.isna().sum().to_dict(),
                    "column_types": {col: str(dtype) for col, dtype in df.dtypes.items()}
                }
            }
            
            # Add validation results if performed
            if validation_results:
                result["validation"] = validation_results
            
            return result
            
        except Exception as e:
            raise FileParsingError(f"Error parsing CSV: {str(e)}")
    
    async def parse_excel(
        self, 
        file_path: Union[str, Path],
        options: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Advanced Excel parsing with multi-sheet support and data validation.
        
        Args:
            file_path: Path to Excel file
            options: Dictionary of options including:
                - sheet_name: Union[str, int, List, None] - Sheets to read
                - header: Union[int, List[str]] - Header row index or custom headers
                - skip_rows: int - Number of rows to skip
                - parse_dates: List[str] - Columns to parse as dates
                - convert_numbers: bool - Auto-convert numeric strings
                - extract_formulas: bool - Extract cell formulas
                - extract_hyperlinks: bool - Extract hyperlinks
                - extract_comments: bool - Extract cell comments
                - validate_schema: Dict - Schema for validation
                
        Returns:
            Dictionary with parsed Excel data and metadata
        """
        if options is None:
            options = {}
        
        try:
            # Prepare pandas read_excel options
            pandas_options = {
                "sheet_name": options.get("sheet_name", None),  # None means all sheets
                "skiprows": options.get("skip_rows", 0),
                "na_values": options.get("na_values", ["", "NA", "N/A", "null", "NULL", "None"])
            }
            
            # Handle custom header
            if "header" in options:
                if isinstance(options["header"], int):
                    pandas_options["header"] = options["header"]
                else:
                    pandas_options["header"] = None
                    pandas_options["names"] = options["header"]
            
            # Parse date columns if specified
            if "parse_dates" in options:
                pandas_options["parse_dates"] = options["parse_dates"]
            
            # Read the Excel file
            excel_data = pd.read_excel(file_path, **pandas_options)
            
            # Process result
            result = {"sheets": {}}
            
            # Handle single sheet or multiple sheets
            if isinstance(excel_data, pd.DataFrame):
                # Single sheet result
                sheet_name = "Sheet1"  # Default name if not specified
                result["sheets"][sheet_name] = self._process_dataframe(
                    excel_data, 
                    options
                )
            else:
                # Multiple sheets
                for sheet_name, df in excel_data.items():
                    result["sheets"][sheet_name] = self._process_dataframe(
                        df, 
                        options
                    )
            
            # Extract additional Excel features if requested
            if any(options.get(opt, False) for opt in ["extract_formulas", "extract_hyperlinks", "extract_comments"]):
                additional_data = self._extract_excel_features(
                    file_path,
                    extract_formulas=options.get("extract_formulas", False),
                    extract_hyperlinks=options.get("extract_hyperlinks", False),
                    extract_comments=options.get("extract_comments", False)
                )
                result["additional_data"] = additional_data
            
            # Add summary information
            result["sheet_count"] = len(result["sheets"])
            result["sheet_names"] = list(result["sheets"].keys())
            
            return result
            
        except Exception as e:
            raise FileParsingError(f"Error parsing Excel file: {str(e)}")
    
    async def parse_txt(
        self, 
        file_path: Union[str, Path],
        options: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Advanced text file parsing with encoding detection and content analysis.
        
        Args:
            file_path: Path to text file
            options: Dictionary of options including:
                - encoding: str - File encoding
                - clean_text: bool - Clean and normalize text
                - extract_entities: bool - Extract named entities
                - summarize: bool - Generate text summary
                - line_numbers: bool - Include line numbers in output
                
        Returns:
            Dictionary with parsed text and analysis
        """
        if options is None:
            options = {}
        
        try:
            # Detect encoding if not specified
            encoding = options.get("encoding", None)
            if not encoding:
                with open(file_path, 'rb') as f:
                    result = chardet.detect(f.read(10000))
                    encoding = result['encoding']
            
            # Read the file
            with open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            
            # Prepare base result
            result = {
                "text": content,
                "character_count": len(content),
                "word_count": len(content.split()),
                "line_count": len(content.splitlines())
            }
            
            # Include lines with line numbers if requested
            if options.get("line_numbers", False):
                lines = content.splitlines()
                result["lines"] = [
                    {"line_number": i+1, "text": line}
                    for i, line in enumerate(lines)
                ]
            
            # Clean text if requested
            if options.get("clean_text", False):
                result["cleaned_text"] = clean_text(content)
            
            # Generate summary if requested
            if options.get("summarize", False):
                result["summary"] = summarize_text(content)
            
            # Extract keywords if requested
            if options.get("extract_keywords", False):
                result["keywords"] = extract_keywords(content)
            
            return result
            
        except Exception as e:
            raise FileParsingError(f"Error parsing text file: {str(e)}")
    
    async def parse_docx(
        self, 
        file_path: Union[str, Path],
        options: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Advanced DOCX parsing with structure and content extraction.
        
        Args:
            file_path: Path to DOCX file
            options: Dictionary of options including:
                - extract_headers: bool - Extract document headers
                - extract_images: bool - Extract embedded images
                - extract_tables: bool - Extract tables
                - include_formatting: bool - Include text formatting info
                
        Returns:
            Dictionary with parsed document structure and content
        """
        if options is None:
            options = {}
        
        try:
            doc = docx.Document(file_path)
            
            # Extract basic content
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():  # Skip empty paragraphs
                    paragraphs.append({
                        "text": para.text,
                        "style": para.style.name
                    })
            
            # Extract document structure
            result = {
                "paragraphs": paragraphs,
                "text": "\n".join(p["text"] for p in paragraphs),
                "metadata": {
                    "title": doc.core_properties.title,
                    "author": doc.core_properties.author,
                    "created": doc.core_properties.created,
                    "modified": doc.core_properties.modified,
                    "last_modified_by": doc.core_properties.last_modified_by,
                    "revision": doc.core_properties.revision
                }
            }
            
            # Extract headers if requested
            if options.get("extract_headers", True):
                headers = []
                for paragraph in doc.paragraphs:
                    if paragraph.style.name.startswith('Heading'):
                        level = int(paragraph.style.name.replace('Heading', ''))
                        headers.append({
                            "level": level,
                            "text": paragraph.text
                        })
                result["headers"] = headers
            
            # Extract tables if requested
            if options.get("extract_tables", True):
                tables = []
                for i, table in enumerate(doc.tables):
                    table_data = []
                    for row in table.rows:
                        row_data = [cell.text for cell in row.cells]
                        table_data.append(row_data)
                    tables.append({
                        "table_index": i,
                        "data": table_data,
                        "row_count": len(table_data),
                        "column_count": len(table_data[0]) if table_data else 0
                    })
                result["tables"] = tables
            
            # TODO: Implement image extraction if requested
            # if options.get("extract_images", False):
            #     # This requires more complex processing with the docx package
            #     pass
            
            return result
            
        except Exception as e:
            raise FileParsingError(f"Error parsing DOCX file: {str(e)}")
    
    async def parse_json(
        self, 
        file_path: Union[str, Path],
        options: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Parse and validate JSON files.
        
        Args:
            file_path: Path to JSON file
            options: Dictionary of options including:
                - encoding: str - File encoding
                - validate_schema: Dict - JSON schema for validation
                - flatten: bool - Flatten nested JSON
                
        Returns:
            Dictionary with parsed JSON and validation results
        """
        if options is None:
            options = {}
        
        try:
            # Detect encoding if not specified
            encoding = options.get("encoding", "utf-8")
            
            # Read and parse JSON
            with open(file_path, 'r', encoding=encoding) as f:
                data = json.load(f)
            
            # Basic analysis of structure
            result = {
                "data": data,
                "structure": self._analyze_json_structure(data)
            }
            
            # TODO: Add schema validation if requested
            # if "validate_schema" in options:
            #     schema = options["validate_schema"]
            #     # Implement JSON schema validation
            
            # TODO: Flatten JSON if requested
            # if options.get("flatten", False):
            #     result["flattened"] = self._flatten_json(data)
            
            return result
            
        except json.JSONDecodeError as e:
            raise FileParsingError(f"Invalid JSON format: {str(e)}")
        except Exception as e:
            raise FileParsingError(f"Error parsing JSON file: {str(e)}")
    
    async def parse_image(
        self, 
        file_path: Union[str, Path],
        options: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Parse images with OCR capabilities and metadata extraction.
        
        Args:
            file_path: Path to image file
            options: Dictionary of options including:
                - ocr: bool - Perform OCR text extraction
                - ocr_lang: str - OCR language
                - extract_exif: bool - Extract EXIF metadata
                - resize: Dict - Resize options
                
        Returns:
            Dictionary with image information and extracted text
        """
        if options is None:
            options = {}
        
        try:
            # Load image
            img = Image.open(file_path)
            
            # Extract basic information
            result = {
                "format": img.format,
                "mode": img.mode,
                "width": img.width,
                "height": img.height,
                "dimensions": f"{img.width}x{img.height}"
            }
            
            # Extract EXIF data if available and requested
            if options.get("extract_exif", True) and hasattr(img, '_getexif') and img._getexif():
                exif_data = {}
                for tag, value in img._getexif().items():
                    if tag in TAGS:
                        exif_data[TAGS[tag]] = value
                result["exif"] = exif_data
            
            # Perform OCR if requested
            if options.get("ocr", False) and FILE_PROCESSING_CONFIG["ocr_enabled"]:
                ocr_lang = options.get("ocr_lang", "eng")
                text = pytesseract.image_to_string(img, lang=ocr_lang)
                result["text"] = text.strip()
                
                # Add text analysis if substantial text found
                if len(text.strip()) > 100:
                    result["text_analysis"] = {
                        "character_count": len(text),
                        "word_count": len(text.split()),
                        "line_count": len(text.splitlines())
                    }
            
            return result
            
        except Exception as e:
            raise FileParsingError(f"Error parsing image: {str(e)}")
    
    def _detect_mime_type(self, file_path: Union[str, Path]) -> str:
        """
        Detect the MIME type of a file.
        
        Args:
            file_path: Path to the file
            
        Returns:
            MIME type string
        """
        try:
            mime = magic.Magic(mime=True)
            return mime.from_file(str(file_path))
        except Exception as e:
            raise FileParsingError(f"Could not detect file type: {str(e)}")
    
    def _determine_file_type(self, mime_type: str) -> str:
        """
        Map MIME type to our supported file types.
        
        Args:
            mime_type: MIME type string
            
        Returns:
            File type category
        """
        for file_type, mime_types in SUPPORTED_FILE_TYPES.items():
            if any(mime in mime_type.lower() for mime in mime_types):
                return file_type
        
        raise UnsupportedFileTypeError(f"Unsupported file type: {mime_type}")
    
    def _get_parsing_method(self, file_type: str) -> Callable:
        """
        Get the appropriate parsing method for a file type.
        
        Args:
            file_type: File type category
            
        Returns:
            Parsing method
        """
        parsing_methods = {
            "pdf": self.parse_pdf,
            "csv": self.parse_csv,
            "excel": self.parse_excel,
            "txt": self.parse_txt,
            "json": self.parse_json,
            "docx": self.parse_docx,
            "image": self.parse_image
        }
        
        if file_type not in parsing_methods:
            raise UnsupportedFileTypeError(f"No parser implemented for file type: {file_type}")
        
        return parsing_methods[file_type]
    
    async def _prepare_input_file(
        self, 
        file_input: Union[str, Path, BinaryIO]
    ) -> Tuple[str, bool]:
        """
        Prepare the input file for processing, handling different input types.
        
        Args:
            file_input: String path, Path object, or file-like object
            
        Returns:
            Tuple of (file_path, should_delete)
        """
        # If it's already a string path
        if isinstance(file_input, str):
            if os.path.isfile(file_input):
                return file_input, False
            else:
                raise FileNotFoundError(f"File not found: {file_input}")
        
        # If it's a Path object
        elif isinstance(file_input, Path):
            if file_input.is_file():
                return str(file_input), False
            else:
                raise FileNotFoundError(f"File not found: {file_input}")
        
        # If it's a file-like object
        elif hasattr(file_input, 'read'):
            # Create temporary file
            temp_file = self.temp_dir / f"temp_{uuid.uuid4()}"
            
            # Write content to temporary file
            with open(temp_file, 'wb') as f:
                if hasattr(file_input, 'seek'):
                    file_input.seek(0)
                content = file_input.read()
                f.write(content if isinstance(content, bytes) else content.encode('utf-8'))
            
            return str(temp_file), True
        
        else:
            raise ValueError(f"Unsupported input type: {type(file_input)}")
    
    def _extract_file_metadata(
        self, 
        file_path: str, 
        file_type: str,
        mime_type: str
    ) -> Dict[str, Any]:
        """
        Extract metadata from file.
        
        Args:
            file_path: Path to file
            file_type: File type category
            mime_type: MIME type
            
        Returns:
            Dictionary of metadata
        """
        metadata = {
            "file_name": os.path.basename(file_path),
            "file_extension": os.path.splitext(file_path)[1],
            "file_size_bytes": os.path.getsize(file_path),
            "creation_time": datetime.fromtimestamp(os.path.getctime(file_path)).isoformat(),
            "modification_time": datetime.fromtimestamp(os.path.getmtime(file_path)).isoformat(),
            "mime_type": mime_type,
            "file_type": file_type
        }
        
        # Add file type specific metadata
        if file_type == "pdf":
            try:
                with pdfplumber.open(file_path) as pdf:
                    metadata.update({
                        "page_count": len(pdf.pages),
                        "pdf_info": pdf.metadata
                    })
            except Exception as e:
                self.processing_stats["warnings"].append(f"Failed to extract PDF metadata: {str(e)}")
        
        elif file_type == "image":
            try:
                with Image.open(file_path) as img:
                    metadata.update({
                        "width": img.width,
                        "height": img.height,
                        "image_format": img.format,
                        "image_mode": img.mode
                    })
            except Exception as e:
                self.processing_stats["warnings"].append(f"Failed to extract image metadata: {str(e)}")
        
        return metadata
    
    def _process_pdf_page(
        self, 
        page, 
        extract_tables: bool,
        ocr_fallback: bool
    ) -> Dict[str, Any]:
        """
        Process a single PDF page.
        
        Args:
            page: pdfplumber page object
            extract_tables: Whether to extract tables
            ocr_fallback: Whether to use OCR if text extraction fails
            
        Returns:
            Dictionary with page content
        """
        page_content = {"text": "", "tables": []}
        
        # Extract text
        text = page.extract_text()
        
        # If text extraction failed and OCR is enabled, try OCR
        if (not text or text.isspace()) and ocr_fallback:
            # Convert page to image
            img = page.to_image()
            img_bytes = img.original.tobytes()
            pil_img = Image.frombytes(
                'RGB', 
                (img.original.width, img.original.height), 
                img_bytes
            )
            
            # Perform OCR
            text = pytesseract.image_to_string(pil_img)
        
        page_content["text"] = text or ""
        
        # Extract tables if requested
        if extract_tables:
            tables = page.extract_tables()
            for i, table in enumerate(tables):
                processed_table = []
                for row in table:
                    processed_row = [
                        str(cell).strip() if cell is not None else "" 
                        for cell in row
                    ]
                    processed_table.append(processed_row)
                
                if processed_table:  # Only add non-empty tables
                    page_content["tables"].append({
                        "table_index": i,
                        "data": processed_table,
                        "row_count": len(processed_table),
                        "column_count": len(processed_table[0]) if processed_table else 0
                    })
        
        return page_content
    
    def _extract_pdf_images(
        self, 
        file_path: str,
        page_range: range
    ) -> List[Dict[str, Any]]:
        """
        Extract images from PDF pages.
        
        Args:
            file_path: Path to PDF
            page_range: Range of pages to process
            
        Returns:
            List of dictionaries with image data
        """
        images = []
        
        try:
            # Convert PDF pages to images
            pdf_images = convert_from_path(
                file_path,
                first_page=min(page_range) + 1,
                last_page=max(page_range) + 1
            )
            
            # Process each image
            for i, img in enumerate(pdf_images):
                page_num = list(page_range)[i]
                image_info = {
                    "page_number": page_num + 1,
                    "width": img.width,
                    "height": img.height,
                    "format": "PNG",
                    "size_bytes": 0  # Will be updated below
                }
                
                # Save to bytes to get size
                img_byte_arr = io.BytesIO()
                img.save(img_byte_arr, format='PNG')
                image_info["size_bytes"] = img_byte_arr.getbuffer().nbytes
                
                # We don't include the actual image data here
                # In a real implementation, we might save it or convert to base64
                images.append(image_info)
                
                # Optionally perform OCR on the image
                # image_info["text"] = pytesseract.image_to_string(img)
        
        except Exception as e:
            self.processing_stats["warnings"].append(f"Failed to extract PDF images: {str(e)}")
        
        return images
    
    def _process_dataframe(
        self, 
        df: pd.DataFrame,
        options: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Process a pandas DataFrame with advanced options.
        
        Args:
            df: Pandas DataFrame
            options: Processing options
            
        Returns:
            Dictionary with processed data
        """
        # Convert numbers if requested
        if options.get("convert_numbers", True):
            for col in df.columns:
                # Try to convert to numeric, but only if most values in the column appear numeric
                if df[col].dtype == 'object':
                    # Sample values to check if they're numbers
                    sample = df[col].dropna().head(100)
                    if sample.empty:
                        continue
                        
                    # Try converting the sample to see if more than 90% convert successfully
                    numeric_count = 0
                    for val in sample:
                        try:
                            float(val)
                            numeric_count += 1
                        except (ValueError, TypeError):
                            pass
                    
                    if numeric_count / len(sample) > 0.9:
                        # Apply fast numeric conversion
                        df[col] = df[col].apply(
                            lambda x: fast_float(x) if x is not None and x != "" else None
                        )
        
        # Validate against schema if provided
        validation_results = {}
        if "validate_schema" in options:
            schema = options["validate_schema"]
            validation_results = validate_data_types(df, schema)
            
            # Validate column names if schema includes them
            if "columns" in schema:
                name_validation = validate_column_names(
                    df.columns.tolist(), 
                    schema["columns"]
                )
                validation_results["column_names"] = name_validation
        
        # Build result
        result = {
            "data": df.to_dict(orient="records"),
            "columns": df.columns.tolist(),
            "row_count": len(df),
            "column_count": len(df.columns),
            "sample": df.head(5).to_dict(orient="records"),
            "summary": {
                "numeric_columns": df.select_dtypes(include=[np.number]).columns.tolist(),
                "date_columns": df.select_dtypes(include=["datetime"]).columns.tolist(),
                "string_columns": df.select_dtypes(include=["object"]).columns.tolist(),
                "missing_values": df.isna().sum().to_dict(),
                "column_types": {col: str(dtype) for col, dtype in df.dtypes.items()}
            }
        }
        
        # Add validation results if performed
        if validation_results:
            result["validation"] = validation_results
        
        return result
    
    def _extract_excel_features(
        self, 
        file_path: str,
        extract_formulas: bool = False,
        extract_hyperlinks: bool = False,
        extract_comments: bool = False
    ) -> Dict[str, Any]:
        """
        Extract advanced Excel features like formulas, hyperlinks and comments.
        
        Args:
            file_path: Path to Excel file
            extract_formulas: Whether to extract formulas
            extract_hyperlinks: Whether to extract hyperlinks
            extract_comments: Whether to extract comments
            
        Returns:
            Dictionary with extracted features
        """
        result = {}
        
        try:
            workbook = load_workbook(file_path, data_only=False)
            
            if extract_formulas:
                formulas = {}
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    sheet_formulas = {}
                    
                    for row in sheet.iter_rows():
                        for cell in row:
                            if cell.value is not None and isinstance(cell.value, str) and cell.value.startswith('='):
                                cell_ref = f"{cell.column_letter}{cell.row}"
                                sheet_formulas[cell_ref] = cell.value
                    
                    if sheet_formulas:
                        formulas[sheet_name] = sheet_formulas
                
                if formulas:
                    result["formulas"] = formulas
            
            if extract_hyperlinks or extract_comments:
                sheet_data = {}
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    sheet_info = {}
                    
                    if extract_hyperlinks:
                        hyperlinks = {}
                        for row in sheet.iter_rows():
                            for cell in row:
                                if cell.hyperlink:
                                    cell_ref = f"{cell.column_letter}{cell.row}"
                                    hyperlinks[cell_ref] = {
                                        "display_text": cell.value,
                                        "url": cell.hyperlink.target
                                    }
                        
                        if hyperlinks:
                            sheet_info["hyperlinks"] = hyperlinks
                    
                    if extract_comments:
                        comments = {}
                        for row in sheet.iter_rows():
                            for cell in row:
                                if cell.comment:
                                    cell_ref = f"{cell.column_letter}{cell.row}"
                                    comments[cell_ref] = str(cell.comment.text)
                        
                        if comments:
                            sheet_info["comments"] = comments
                    
                    if sheet_info:
                        sheet_data[sheet_name] = sheet_info
                
                if sheet_data:
                    result["sheet_data"] = sheet_data
        
        except Exception as e:
            self.processing_stats["warnings"].append(f"Failed to extract Excel features: {str(e)}")
        
        return result
    
    def _analyze_json_structure(self, data: Any, max_depth: int = 5) -> Dict[str, Any]:
        """
        Analyze the structure of a JSON document.
        
        Args:
            data: JSON data
            max_depth: Maximum recursion depth
            
        Returns:
            Dictionary with structure analysis
        """
        if max_depth <= 0:
            return {"type": "max_depth_reached"}
        
        if isinstance(data, dict):
            return {
                "type": "object",
                "properties": {
                    k: self._analyze_json_structure(v, max_depth - 1) 
                    for k, v in list(data.items())[:100]  # Limit to first 100 keys for large objects
                },
                "property_count": len(data)
            }
        elif isinstance(data, list):
            if not data:
                return {"type": "array", "items_type": "empty", "length": 0}
            
            # Sample items for large arrays
            sample_size = min(10, len(data))
            samples = data[:sample_size]
            
            # Check if all items have the same type
            item_types = set(type(item).__name__ for item in samples)
            
            if len(item_types) == 1:
                # All items are the same type
                return {
                    "type": "array",
                    "items_type": next(iter(item_types)),
                    "length": len(data),
                    "sample": self._analyze_json_structure(samples[0], max_depth - 1) if samples else None
                }
            else:
                # Mixed types
                return {
                    "type": "array",
                    "items_type": "mixed",
                    "type_distribution": {t: sum(1 for item in samples if type(item).__name__ == t) for t in item_types},
                    "length": len(data)
                }
        else:
            # Primitive types
            return {"type": type(data).__name__}
    
    async def _record_processing_log(
        self, 
        processing_id: str,
        success: bool,
        details: Dict[str, Any]
    ) -> None:
        """
        Record processing log to database.
        
        Args:
            processing_id: Unique processing ID
            success: Whether processing was successful
            details: Processing details
        """
        if not self.db:
            return
        
        try:
            # Create log record
            log = ProcessingLog(
                processing_id=processing_id,
                success=success,
                processing_time=details.get("duration_seconds", 0),
                details=details,
                created_at=datetime.utcnow()
            )
            
            # Add to database
            self.db.add(log)
            await self.db.commit()
            
        except Exception as e:
            file_logger.error(f"Failed to record processing log: {str(e)}")


# Factory function to get file parser instance
async def get_file_parser(db: Optional[AsyncSession] = None) -> EnterpriseFileParser:
    """
    Get a configured file parser instance.
    
    Args:
        db: Optional database session
        
    Returns:
        EnterpriseFileParser instance
    """
    return EnterpriseFileParser(db)


# Convenience functions for direct use

async def parse_file(
    file_path: Union[str, Path, BinaryIO],
    file_type: Optional[str] = None,
    options: Optional[Dict[str, Any]] = None,
    db: Optional[AsyncSession] = None
) -> Dict[str, Any]:
    """
    Parse a file with automatic type detection and comprehensive processing.
    
    Args:
        file_path: Path to file, Path object, or file-like object
        file_type: Optional explicit file type override
        options: Optional processing parameters
        db: Optional database session for logging
        
    Returns:
        Dictionary with parsed content and metadata
    """
    parser = await get_file_parser(db)
    return await parser.parse_file(file_path, file_type, options)


async def parse_pdf(
    file_path: Union[str, Path],
    options: Optional[Dict[str, Any]] = None,
    db: Optional[AsyncSession] = None
) -> Dict[str, Any]:
    """
    Parse a PDF file with advanced options.
    
    Args:
        file_path: Path to PDF file
        options: Processing options
        db: Optional database session for logging
        
    Returns:
        Dictionary with PDF content
    """
    parser = await get_file_parser(db)
    return await parser.parse_pdf(file_path, options)


async def parse_csv(
    file_path: Union[str, Path],
    options: Optional[Dict[str, Any]] = None,
    db: Optional[AsyncSession] = None
) -> Dict[str, Any]:
    """
    Parse a CSV file with advanced options.
    
    Args:
        file_path: Path to CSV file
        options: Processing options
        db: Optional database session for logging
        
    Returns:
        Dictionary with CSV data
    """
    parser = await get_file_parser(db)
    return await parser.parse_csv(file_path, options)


async def auto_detect_file_type(file_path: Union[str, Path]) -> str:
    """
    Detect the MIME type of a file.
    
    Args:
        file_path: Path to file
        
    Returns:
        MIME type string
    """
    try:
        mime = magic.Magic(mime=True)
        return mime.from_file(str(file_path))
    except Exception as e:
        raise FileParsingError(f"Could not detect file type: {str(e)}")